import os
import sys
import argparse
from pathlib import Path
import time
import re
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from typing import List, Dict
import ollama
from docx import Document

def process_file(file_path_str: str, model_name: str) -> Dict:
    """Process a single .docx file and return its content and generated essay with Markdown formatting."""
    file_path = Path(file_path_str)
    try:
        client = ollama.Client()

        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables_content = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                tables_content.append("\n".join(table_text))

        content = {
            "paragraphs": "\n\n".join(paragraphs),
            "tables": "\n\n".join(tables_content) if tables_content else "",
            "word_count": len(" ".join(paragraphs).split()),
            "paragraph_count": len(paragraphs),
        }

        if not content['paragraphs']:
            essay_text = f"No readable content found in {file_path.name}"
        else:
            prompt = f"""
Please write a comprehensive analytical essay about the document "{file_path.name}" with the following structure, formatted in Markdown:

## Document Overview
Briefly describe the document's purpose and content.

## Key Themes and Topics
List and describe key themes and topics identified.

## Writing Style and Structure Analysis
Analyse the document's writing style and structure.

## Main Arguments or Points Presented
Summarise the core arguments or points.

## Critical Assessment
Provide a critical assessment.

## Conclusions and Significance
Summarise the document's significance and final thoughts.

**Document Statistics:**
- Word count: {content['word_count']}
- Paragraph count: {content['paragraph_count']}
- File path: {file_path}

**Document Content Preview:**
{content['paragraphs'][:500]}...

{f"**Tables/Structured Data:**\n{content['tables']}" if content['tables'] else ""}
            """
            response = client.generate(model=model_name, prompt=prompt)
            essay_text = response['response']

        result = {
            'file_path': file_path_str,
            'content': content,
            'essay': essay_text
        }
        return result

    except Exception as e:
        return {
            'file_path': file_path_str,
            'error': str(e),
            'content': {},
            'essay': ""
        }

class DocxAnalyzer:
    def __init__(self, model_name: str = "llama3.2"):
        self.model_name = model_name
        self.file_counter = 0

    def find_docx_files(self, directory: str) -> List[Path]:
        directory_path = Path(directory)
        return [p for p in directory_path.rglob("*.docx") if not p.name.startswith("~$")]

    def analyze_directory(self, directory: str, output_dir: str = None):
        docx_files = self.find_docx_files(directory)
        if not docx_files:
            print(f"No .docx files found in '{directory}'")
            return

        if output_dir is None:
            folder_slug = re.sub(r'[^a-zA-Z0-9]+', '_', Path(directory).name.lower())
            script_root = Path(__file__).parent
            output_dir = script_root / f"{folder_slug}_essays"
        else:
            output_dir = Path(output_dir)

        output_dir.mkdir(exist_ok=True)

        use_multiprocessing = os.cpu_count() >= 4
        Executor = ProcessPoolExecutor if use_multiprocessing else ThreadPoolExecutor
        max_workers = min(os.cpu_count() or 2, 8)

        print(f"Using {'multiprocessing' if use_multiprocessing else 'threading'} with {max_workers} workers...")

        all_results = []

        with Executor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(process_file, str(file_path), self.model_name): file_path
                for file_path in docx_files
            }

            for future in as_completed(futures):
                result = future.result()
                all_results.append(result)

                self.file_counter += 1
                file_path = Path(result['file_path'])
                markdown_filename = f"{self.file_counter:02d}_{file_path.stem}_analysis.md"
                markdown_path = output_dir / markdown_filename

                if "error" in result:
                    markdown_content = f"# Error: {result['error']}\n\nFile: {file_path}\n"
                else:
                    metadata = (
                        f"# Document Analysis for {file_path.name}\n\n"
                        f"**Analysis Date:** {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
                        f"**Word Count:** {result['content'].get('word_count', 'N/A')}\n"
                        f"**Paragraph Count:** {result['content'].get('paragraph_count', 'N/A')}\n\n"
                        "---\n\n"
                    )
                    markdown_content = metadata + result['essay'] + f"\n\n---\n\n*Generated using Ollama model: {self.model_name}*"

                with open(markdown_path, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)

                print(f"Processed {file_path.name}")

        print(f"✓ Analysis complete! Markdown essays saved to: {output_dir}")

def main():
    parser = argparse.ArgumentParser(description="Analyse .docx files and generate Markdown essays using Ollama")
    parser.add_argument("directory", help="Directory to search for .docx files")
    parser.add_argument("-o", "--output", help="Output directory for essays")
    parser.add_argument("-m", "--model", default="llama3.2", help="Ollama model to use (default: llama3.2)")

    args = parser.parse_args()

    if not os.path.exists(args.directory):
        print(f"Error: Directory '{args.directory}' does not exist")
        sys.exit(1)

    analyzer = DocxAnalyzer(model_name=args.model)

    try:
        analyzer.analyze_directory(args.directory, args.output)
    except KeyboardInterrupt:
        print("\nAnalysis interrupted by user")
    except Exception as e:
        print(f"Error during analysis: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
